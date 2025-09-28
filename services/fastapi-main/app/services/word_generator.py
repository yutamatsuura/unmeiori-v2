"""
Word文書生成サービス
印刷プレビューと同等のレイアウト・区切り線を再現
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from io import BytesIO
from typing import Dict, Any, List, Optional
from datetime import datetime


class WordDocumentGenerator:
    """Word文書生成クラス"""

    # 色定義（印刷プレビューCSS準拠）
    PRIMARY_COLOR = RGBColor(25, 118, 210)      # #1976d2
    SECONDARY_COLOR = RGBColor(220, 0, 78)      # #dc004e
    BORDER_COLOR = RGBColor(224, 224, 224)      # #e0e0e0
    TEXT_DARK = RGBColor(66, 66, 66)            # #424242
    TEXT_GRAY = RGBColor(117, 117, 117)         # #757575

    def __init__(self):
        self.doc = Document()
        self._setup_document_styles()

    def _setup_document_styles(self):
        """ドキュメント基本スタイル設定"""
        # ページ設定（A4）
        section = self.doc.sections[0]
        section.page_height = Inches(11.69)  # A4 height
        section.page_width = Inches(8.27)    # A4 width
        section.left_margin = Inches(0.59)   # 15mm
        section.right_margin = Inches(0.59)  # 15mm
        section.top_margin = Inches(0.59)    # 15mm
        section.bottom_margin = Inches(0.59) # 15mm

        # デフォルトフォント設定
        style = self.doc.styles['Normal']
        style.font.name = 'Yu Gothic'
        style.font.size = Pt(11)
        style.font.color.rgb = self.TEXT_DARK

    def add_header_with_divider(self, text: str, level: int = 1):
        """ヘッダーを追加"""
        if level == 1:
            heading = self.doc.add_heading(text, level=1)
            heading.runs[0].font.size = Pt(24)
            heading.runs[0].font.color.rgb = self.PRIMARY_COLOR
        else:
            heading = self.doc.add_heading(text, level=level)
            heading.runs[0].font.size = Pt(16)
            heading.runs[0].font.color.rgb = self.PRIMARY_COLOR

        heading.paragraph_format.space_after = Pt(12)
        return heading

    def add_section_title_with_divider(self, text: str):
        """セクションタイトルを追加"""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = self.PRIMARY_COLOR

        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.space_before = Pt(12)

        return para

    def add_info_row_with_divider(self, label: str, value: str):
        """情報行を追加"""
        para = self.doc.add_paragraph()

        # ラベル（太字）
        label_run = para.add_run(f"{label}: ")
        label_run.font.bold = True
        label_run.font.color.rgb = self.TEXT_GRAY
        label_run.font.size = Pt(11)

        # 値
        value_run = para.add_run(value)
        value_run.font.color.rgb = self.TEXT_DARK
        value_run.font.size = Pt(11)

        para.paragraph_format.space_after = Pt(4)

        return para

    def create_client_info_table(self, form_data: Dict[str, Any]):
        """基本情報テーブル作成"""
        table = self.doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'

        # 1行目
        cells = table.rows[0].cells
        cells[0].text = "氏名:"
        cells[1].text = form_data.get('name', '')
        cells[2].text = "性別:"
        cells[3].text = "男性" if form_data.get('gender') == 'male' else "女性"

        # 2行目
        cells = table.rows[1].cells
        cells[0].text = "生年月日:"
        cells[1].text = form_data.get('birthDate', '')
        cells[2].text = "鑑定日:"
        cells[3].text = datetime.now().strftime('%Y年%m月%d日')

        # テーブルスタイル設定
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        if ":" in run.text:
                            run.font.bold = True
                            run.font.color.rgb = self.TEXT_GRAY
                        else:
                            run.font.color.rgb = self.TEXT_DARK

        return table

    def create_results_grid(self, items: List[Dict[str, str]], cols: int = 3):
        """結果グリッド作成（九星気学・姓名判断用）"""
        rows = (len(items) + cols - 1) // cols  # 必要行数計算
        table = self.doc.add_table(rows=rows * 2, cols=cols)  # ラベル行+値行

        for i, item in enumerate(items):
            col = i % cols
            row_base = (i // cols) * 2

            # ラベル行
            label_cell = table.cell(row_base, col)
            label_cell.text = item['label']
            label_para = label_cell.paragraphs[0]
            label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label_run = label_para.runs[0]
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = self.TEXT_GRAY

            # 値行
            value_cell = table.cell(row_base + 1, col)
            value_cell.text = item['value']
            value_para = value_cell.paragraphs[0]
            value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            value_run = value_para.runs[0]
            value_run.font.size = Pt(12)
            value_run.font.bold = True
            value_run.font.color.rgb = self.TEXT_DARK

        # テーブル全体のスタイル
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # セル境界線設定
                cell._element.get_or_add_tcPr().append(
                    parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                             f'<w:top w:val="single" w:sz="8" w:color="e0e0e0"/>'
                             f'<w:bottom w:val="single" w:sz="8" w:color="e0e0e0"/>'
                             f'<w:left w:val="single" w:sz="8" w:color="e0e0e0"/>'
                             f'<w:right w:val="single" w:sz="8" w:color="e0e0e0"/>'
                             f'</w:tcBorders>')
                )

        return table

    def create_houi_details_table(self, houi_details: List[Dict[str, Any]]):
        """方位詳細テーブル作成"""
        table = self.doc.add_table(rows=len(houi_details) + 1, cols=5)
        table.style = 'Table Grid'

        # ヘッダー行
        header_cells = table.rows[0].cells
        headers = ['方位', '年盤', '月盤', '日盤', '判定']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_para = header_cells[i].paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_run = header_para.runs[0]
            header_run.font.bold = True
            header_run.font.size = Pt(11)
            header_run.font.color.rgb = self.TEXT_DARK

            # ヘッダー背景色
            header_cells[i]._element.get_or_add_tcPr().append(
                parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="f5f5f5"/>')
            )

        # データ行
        for i, detail in enumerate(houi_details):
            row = table.rows[i + 1]
            cells = row.cells

            cells[0].text = detail.get('houi', '')
            cells[1].text = str(detail.get('nenbanStar', ''))
            cells[2].text = str(detail.get('getsubanStar', ''))
            cells[3].text = str(detail.get('nippanStar', '')) if detail.get('nippanStar') else '-'

            # 判定セル（色分け）
            judgment_cell = cells[4]
            if detail.get('kyouType'):
                judgment_cell.text = detail['kyouType']
                # 赤系背景色
                judgment_cell._element.get_or_add_tcPr().append(
                    parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="ffe6e6"/>')
                )
            elif detail.get('kichiType'):
                judgment_cell.text = detail['kichiType']
                # 青系背景色
                judgment_cell._element.get_or_add_tcPr().append(
                    parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="e6f3ff"/>')
                )
            else:
                judgment_cell.text = '-'

            # セル中央揃え
            for cell in cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        return table

    def add_hoiban_image(self, image_bytes: bytes):
        """方位盤画像追加"""
        try:
            self.doc.add_picture(BytesIO(image_bytes), width=Inches(4))
            # 中央揃え
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            # 画像追加失敗時はプレースホルダーテキスト
            para = self.doc.add_paragraph("【方位盤画像】")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.font.size = Pt(12)
            run.font.italic = True
            run.font.color.rgb = self.TEXT_GRAY

    def add_footer(self, target_date: str):
        """フッター追加"""
        section = self.doc.sections[0]
        footer = section.footer

        footer_table = footer.add_table(1, 2, Inches(6))

        # 日付（左）
        date_cell = footer_table.cell(0, 0)
        date_cell.text = target_date
        date_para = date_cell.paragraphs[0]
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_run = date_para.runs[0]
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = self.TEXT_GRAY

        # システム名（右）
        system_cell = footer_table.cell(0, 1)
        system_cell.text = "運命織 - プロフェッショナル鑑定書作成システム"
        system_para = system_cell.paragraphs[0]
        system_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        system_run = system_para.runs[0]
        system_run.font.size = Pt(9)
        system_run.font.color.rgb = self.TEXT_GRAY

    def generate_document(self, data: Dict[str, Any]) -> bytes:
        """Word文書生成メイン処理"""
        try:
            # 1. ヘッダー
            self.add_header_with_divider("総合鑑定書", level=1)
            subtitle_para = self.doc.add_paragraph("九星気学・姓名判断による詳細鑑定")
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.runs[0]
            subtitle_run.font.size = Pt(12)
            subtitle_run.font.color.rgb = self.TEXT_GRAY

            subtitle_para.paragraph_format.space_after = Pt(16)

            # 2. 基本情報
            self.create_client_info_table(data.get('formData', {}))
            self.doc.add_paragraph()  # スペース

            # 3. 九星気学結果
            self.add_section_title_with_divider("九星気学結果")
            if data.get('kyuseiResult'):
                kyusei_items = [
                    {'label': '本命星', 'value': data['kyuseiResult'].get('honmei', '未計算')},
                    {'label': '月命星', 'value': data['kyuseiResult'].get('getsusei', '未計算')},
                    {'label': '日命星', 'value': data['kyuseiResult'].get('nissei', '未計算')}
                ]
                self.create_results_grid(kyusei_items, cols=3)

            # 4. 方位盤
            if data.get('hoibanImage'):
                self.add_hoiban_image(data['hoibanImage'])

            # 5. 姓名判断結果
            self.add_section_title_with_divider("姓名判断結果")
            if data.get('seimeiResult'):
                seimei_items = [
                    {'label': '天格', 'value': f"{data['seimeiResult'].get('heaven', '未計算')}画"},
                    {'label': '人格', 'value': f"{data['seimeiResult'].get('personality', '未計算')}画"},
                    {'label': '地格', 'value': f"{data['seimeiResult'].get('earth', '未計算')}画"},
                    {'label': '総画', 'value': f"{data['seimeiResult'].get('total', '未計算')}画"}
                ]
                self.create_results_grid(seimei_items, cols=4)

            # 6. 詳細情報
            self.add_section_title_with_divider("詳細情報")

            # 基本情報
            if data.get('birthData'):
                birth_data = data['birthData']
                if birth_data.get('year'):
                    self.add_info_row_with_divider("本命星", birth_data['year'].get('name', ''))
                if birth_data.get('month'):
                    self.add_info_row_with_divider("月命星", birth_data['month'].get('name', ''))
                if birth_data.get('day'):
                    self.add_info_row_with_divider("日命星", birth_data['day'].get('name', ''))

            # 7. 方位詳細
            if data.get('houiDetails'):
                self.doc.add_paragraph()  # スペース
                self.add_section_title_with_divider("方位詳細")
                self.create_houi_details_table(data['houiDetails'])

            # 8. 鑑定士コメント
            if data.get('kantei_comment'):
                self.doc.add_paragraph()  # スペース
                self.add_section_title_with_divider("鑑定士からのコメント")

                # 改行を考慮してコメントを追加
                comment_text = data['kantei_comment']
                comment_lines = comment_text.split('\n')

                for i, line in enumerate(comment_lines):
                    if i == 0:
                        # 最初の行は新しい段落として
                        comment_para = self.doc.add_paragraph(line)
                    else:
                        # 2行目以降は改行として追加
                        comment_para.add_run('\n' + line)

                comment_para.paragraph_format.space_after = Pt(16)
                for run in comment_para.runs:
                    run.font.size = Pt(11)
                    run.font.color.rgb = self.TEXT_DARK

            # 9. フッター
            self.add_footer(data.get('targetDate', datetime.now().strftime('%Y年%m月%d日')))

            # Word文書をバイト列として返却
            buffer = BytesIO()
            self.doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()

        except Exception as e:
            raise Exception(f"Word文書生成エラー: {str(e)}")


def generate_word_document(data: Dict[str, Any]) -> bytes:
    """Word文書生成関数（外部API用）"""
    generator = WordDocumentGenerator()
    return generator.generate_document(data)